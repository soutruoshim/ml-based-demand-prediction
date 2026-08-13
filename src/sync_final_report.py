#!/usr/bin/env python3
"""Synchronize the paper's tables, narrative values, sample rows, and charts."""
from pathlib import Path
from tempfile import NamedTemporaryFile
from zipfile import ZIP_DEFLATED, ZipFile

import pandas as pd
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "Final_Deliverables_ML_Demand_Prediction.docx"
OUT = ROOT / "outputs"
DATA = ROOT / "data" / "synthetic_factory_demand_dataset.csv"


def replace_all(document, replacements):
    containers = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                containers.extend(cell.paragraphs)
    for paragraph in containers:
        original = paragraph.text
        updated = original
        for old, new in replacements.items():
            updated = updated.replace(old, new)
        if updated != original:
            paragraph.text = updated


def main():
    evaluation = pd.read_csv(OUT / "model_evaluation_results.csv")
    importance = pd.read_csv(OUT / "feature_importance.csv")
    plan = pd.read_csv(OUT / "jan_2026_production_recommendation.csv")
    data = pd.read_csv(DATA)
    best = evaluation.iloc[0]
    doc = Document(REPORT)

    replacements = {
        "MAE 54.1, RMSE 66.4, and R2 0.961": f"MAE {best.MAE:.1f}, RMSE {best.RMSE:.1f}, and R2 {best.R2:.3f}",
        "MAE 54.10, RMSE 66.40, and R2 0.9607": f"MAE {best.MAE:.2f}, RMSE {best.RMSE:.2f}, and R2 {best.R2:.4f}",
        "best-model R2 of 0.961": f"best-model R2 of {best.R2:.3f}",
    }
    replace_all(doc, replacements)

    # Model evaluation table.
    table = doc.tables[5]
    lookup = {r.Model: r for r in evaluation.itertuples()}
    for row in table.rows[1:]:
        r = lookup[row.cells[0].text]
        row.cells[1].text, row.cells[2].text, row.cells[3].text = f"{r.MAE:.2f}", f"{r.RMSE:.2f}", f"{r.R2:.4f}"

    # Feature-importance table: show the six strongest predictors.
    table = doc.tables[6]
    for row, r in zip(table.rows[1:], importance.head(6).itertuples()):
        row.cells[0].text, row.cells[1].text = r.Feature, f"{getattr(r, '_2'):.2f}"

    # January planning table.
    table = doc.tables[7]
    lookup = {r.product_type: r for r in plan.itertuples()}
    for row in table.rows[1:]:
        r = lookup[row.cells[0].text]
        vals = [r.predicted_demand, r.stock_quantity, r.safety_stock_10_percent, r.recommended_production]
        for cell, value in zip(row.cells[1:], vals): cell.text = f"{value:,.0f}"

    # Appendix sample rows must be actual rows from the submitted CSV.
    table = doc.tables[9]
    for row, r in zip(table.rows[1:], data.head(9).itertuples()):
        vals = [r.forecast_month, r.product_type, r.previous_sales, r.stock_quantity, f"{r.price:.2f}", r.promotion, r.customer_order_quantity, r.target_demand]
        for cell, value in zip(row.cells, vals): cell.text = str(value)

    doc.save(REPORT)

    # The report has one workflow image followed by the four generated result figures.
    charts = {
        "word/media/image2.png": OUT / "figures" / "model_rmse.png",
        "word/media/image3.png": OUT / "figures" / "actual_vs_predicted.png",
        "word/media/image4.png": OUT / "figures" / "demand_trends.png",
        "word/media/image5.png": OUT / "figures" / "feature_importance.png",
    }
    with NamedTemporaryFile(suffix=".docx", dir=ROOT, delete=False) as tmp:
        temp_path = Path(tmp.name)
    with ZipFile(REPORT, "r") as source, ZipFile(temp_path, "w", ZIP_DEFLATED) as target:
        for info in source.infolist():
            payload = charts[info.filename].read_bytes() if info.filename in charts else source.read(info.filename)
            target.writestr(info, payload)
    temp_path.replace(REPORT)
    print(f"Synchronized {REPORT.name}")


if __name__ == "__main__": main()
